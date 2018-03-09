#! python3
# customInvitations.py - generate a Word document with custom invitations 

import docx, os


readFile = open('guests.txt')
dates = readFile.readlines()
readFile.close()

doc = docx.Document('invitations.docx')

for date in dates:
    doc.add_paragraph('It would be a pleasure to have the conpany of')
    doc.add_heading(date[:-1], 0)
    doc.add_paragraph('at 11010 Memary Lane on the Evening of')
    doc.add_heading('April 1st', 2)
    doc.add_paragraph('at 7 o\'clock')
    
    doc.paragraphs[len(doc.paragraphs)-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

doc.save('invitations.docx')